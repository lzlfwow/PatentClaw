from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pytest
from PIL import Image
from docx import Document
from docx.oxml.ns import qn

from latex2disclosure.config import Settings
from latex2disclosure.agents.invention_mining import _patent_title
from latex2disclosure.pipeline import DisclosurePipeline
from latex2disclosure.schemas import JobStatus
from latex2disclosure.storage import JobStore


def test_patent_title_compacts_without_cutting_the_type_suffix():
    title = "一种视觉语言模型的语言侧推理向量提取、迁移与干预方法及系统"
    result = _patent_title(title)
    assert result == "一种视觉语言模型语言侧推理向量干预方法及系统"
    assert len(result) <= 25


@pytest.mark.asyncio
async def test_offline_pipeline_generates_reviewable_disclosure(tmp_path: Path):
    original_sample = Path(__file__).parents[1] / "examples" / "sample.tex"
    sample = tmp_path / "sample.tex"
    sample.write_text(
        original_sample.read_text(encoding="utf-8").replace(
            "\\begin{figure}\n",
            "\\begin{figure}\n\\includegraphics{paper-original.png}\n",
        ),
        encoding="utf-8",
    )
    source_figure = tmp_path / "paper-original.png"
    Image.new("RGB", (120, 80), color=(210, 30, 30)).save(source_figure)
    settings = Settings(L2D_OFFLINE_MODE=True, L2D_DATA_DIR=tmp_path / "data").prepare()
    store = JobStore(settings.data_dir / "jobs")
    pipeline = DisclosurePipeline(settings, store)
    job = pipeline.create(sample.name, sample)
    completed = await pipeline.run(job.job_id)

    assert completed.status == JobStatus.completed
    assert completed.paper is not None
    assert completed.understanding is not None
    assert completed.invention is not None
    assert completed.solution is not None
    assert completed.evidence_package is not None
    assert len(completed.evidence_package.patent_figures) >= 1
    assert completed.disclosure is not None
    assert completed.review is not None
    assert completed.review.passed
    assert completed.metadata["figure_source_policy"] == "regenerated_only"
    assert len(completed.invention.proposed_title) <= 25
    assert len(completed.solution.method_steps) >= 3
    assert set(completed.artifacts) == {"markdown", "docx", "json", "figures"}
    assert [event.status for event in completed.events].count("completed") == 8

    snapshot = json.loads(Path(completed.artifacts["json"]).read_text(encoding="utf-8"))
    assert snapshot["status"] == "completed"
    assert snapshot["events"][-1]["stage"] == "stage7_export"
    assert snapshot["events"][-1]["status"] == "completed"
    with zipfile.ZipFile(completed.artifacts["figures"]) as archive:
        names = set(archive.namelist())
        assert "figure-1.png" in names
        assert "figure-1.mmd" in names
        assert "manifest.json" in names
        assert "未复用论文原图" in archive.read("figure-1.mmd").decode("utf-8")
    document = Document(completed.artifacts["docx"])
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert document.paragraphs[0].text == "技　术　交　底　书"
    assert "1. 发明创造名称" in text
    assert "2. 背景技术及现有技术的缺陷和不足" in text
    assert "3. 具体的技术方案描述" in text
    assert "4. 本发明创造的优点" in text
    assert "5. 具体实施方式及附图" in text
    assert "详细技术方案" in text
    assert "系统实现" in text
    assert "系统实现与数据接口" in text
    assert "术语、实施边界与替代方案" in text
    assert "项目编号" not in text
    assert all(not paragraph.text for paragraph in document.sections[0].header.paragraphs)
    assert all(not paragraph.text for paragraph in document.sections[0].footer.paragraphs)
    normal = document.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal._element.rPr.rFonts.get(qn("w:eastAsia")) == "宋体"
    assert normal.font.size.pt == pytest.approx(12)
    heading = next(paragraph for paragraph in document.paragraphs if paragraph.style.name == "Heading 1")
    assert heading.runs[0].font.size.pt == pytest.approx(14)
    assert heading.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
    assert str(heading.runs[0].font.color.rgb) == "000000"
    assert len(document.inline_shapes) == len(completed.evidence_package.patent_figures)
    with zipfile.ZipFile(completed.artifacts["docx"]) as archive:
        embedded_images = [archive.read(name) for name in archive.namelist() if name.startswith("word/media/")]
    assert embedded_images
    assert source_figure.read_bytes() not in embedded_images
