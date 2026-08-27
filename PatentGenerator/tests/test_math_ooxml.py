from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from latex2disclosure.math_ooxml import append_rich_text, normalize_math_markup


def test_legacy_math_is_normalized_and_exported_as_editable_omml():
    text = (
        "计算δ_l(a)=h_l(a;C_R)-h_l(a;C_S)，并执行"
        "x̃_l=x_l+μr_l，学习率为1×10^-4。"
    )
    normalized = normalize_math_markup(text)
    assert "$\\boldsymbol{\\delta}_l(a)=" in normalized
    assert "$1\\times10^{-4}$" in normalized

    document = Document()
    paragraph = document.add_paragraph()
    append_rich_text(paragraph, text)
    xml = paragraph._p.xml
    assert "m:oMath" in xml
    assert "m:sSub" in xml
    assert "m:sSup" in xml
    assert "x̃_l" not in xml
    assert "10^-4" not in xml


def test_calligraphic_and_vector_styles_are_preserved():
    document = Document()
    paragraph = document.add_paragraph()
    append_rich_text(paragraph, "$\\mathcal{C}_R=(Q,T,A)$ and $\\boldsymbol{r}_l$")
    runs = paragraph._p.findall(f".//{qn('m:r')}")
    assert runs
    assert paragraph._p.findall(f".//{qn('m:sSub')}")
    assert paragraph._p.findall(f".//{qn('m:scr')}")
    assert paragraph._p.findall(f".//{qn('w:b')}")


def test_math_runs_define_all_word_font_slots():
    document = Document()
    paragraph = document.add_paragraph()
    append_rich_text(paragraph, "$\\boldsymbol{\\delta}_l(a)$")

    fonts = paragraph._p.find(f".//{qn('w:rFonts')}")
    assert fonts is not None
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        assert fonts.get(qn(f"w:{slot}")) == "Cambria Math"
    assert paragraph._p.find(f".//{qn('w:noProof')}") is not None
