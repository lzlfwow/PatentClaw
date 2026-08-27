from __future__ import annotations

import re
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from .patent_figures import (
    ExperimentSeries,
    build_default_patent_figures,
    drawing_descriptions,
    extract_experiment_series,
    render_experiment_chart,
    render_patent_figures,
)
from .math_ooxml import append_rich_text, display_math_element, normalize_math_markup
from .schemas import PipelineJob


BLACK = RGBColor(0, 0, 0)
LATIN_FONT = "Times New Roman"
BODY_CN_FONT = "宋体"
HEADING_CN_FONT = "黑体"


def _join(items: list[str], fallback: str = "待补充") -> str:
    return "\n".join(items) if items else fallback


def markdown(job: PipelineJob) -> str:
    disclosure = job.disclosure
    if disclosure is None:
        raise ValueError("技术交底书尚未生成")
    lines = [
        f"# {disclosure.invention_title}",
        "",
        f"> 项目编号：{job.job_id}｜来源：{job.input_name}",
        "",
        "## 1. 技术领域",
        "",
        disclosure.technical_field,
        "",
        "## 2. 背景技术",
        "",
        *[f"- {item}" for item in disclosure.background],
        "",
        "## 3. 现有技术缺陷与技术问题",
        "",
        *[f"- {item}" for item in disclosure.prior_art_defects + disclosure.technical_problem],
        "",
        "## 4. 技术方案",
        "",
        disclosure.overall_solution,
        "",
        *[f"{index}. {item}" for index, item in enumerate(disclosure.detailed_steps, start=1)],
        "",
        "## 5. 关键创新点",
        "",
        *[f"- {item}" for item in disclosure.key_innovations],
        "",
        "## 6. 有益效果",
        "",
        *[f"- {item}" for item in disclosure.beneficial_effects],
        "",
        "## 7. 具体实施方式",
        "",
        *[f"### 实施例{index}\n\n{item}" for index, item in enumerate(disclosure.embodiments, start=1)],
        "",
        "## 8. 实验依据",
        "",
        *[f"- {item}" for item in disclosure.experimental_evidence],
        "",
        "## 9. 附图说明",
        "",
        *[f"- {item}" for item in disclosure.drawing_descriptions],
        "",
        "## 10. 系统实现",
        "",
        *[f"- {item}" for item in disclosure.system_implementation],
        "",
        "## 11. 数据与接口",
        "",
        *[f"- {item}" for item in disclosure.data_and_interfaces],
        "",
        "## 12. 术语与符号",
        "",
        *[f"- {item}" for item in disclosure.terminology],
        "",
        "## 13. 实施边界",
        "",
        *[f"- {item}" for item in disclosure.implementation_boundaries],
        "",
        "## 14. 替代实施方式",
        "",
        *[f"- {item}" for item in disclosure.alternatives],
        "",
        "## 15. 发明人确认事项",
        "",
        *[f"- {item}" for item in disclosure.inventor_confirmation_items],
    ]
    if job.evidence_package and job.evidence_package.patent_figures:
        lines.extend(["", "## 专利附图", ""])
        for figure in job.evidence_package.patent_figures:
            lines.extend(
                [
                    f"![图{figure.figure_no} {figure.title}](figures/figure-{figure.figure_no}.png)",
                    "",
                    f"图{figure.figure_no}  {figure.title}",
                    "",
                ]
            )
        lines.append("> 附图由技术方案重新生成，未复制或嵌入论文原始图片；可编辑源文件见附图包。")
    lines.extend(["", "---", "本文件由Agent辅助生成，仅供发明人和专利代理师审阅，不构成法律意见。"])
    return normalize_math_markup("\n".join(lines))


def _set_font(
    run,
    size: float,
    bold: bool = False,
    east_asia_font: str = BODY_CN_FONT,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = BLACK


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    for name, size, before, after, east_asia_font in (
        ("Heading 1", 14, 15, 8, HEADING_CN_FONT),
        ("Heading 2", 12, 11, 6, HEADING_CN_FONT),
        ("Heading 3", 12, 8, 4, BODY_CN_FONT),
    ):
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.keep_with_next = True


def _add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(level=level)
    _set_font(
        paragraph.add_run(text),
        14 if level == 1 else 12,
        bold=True,
        east_asia_font=HEADING_CN_FONT if level < 3 else BODY_CN_FONT,
    )


def _add_body(document: Document, text: str, *, bold: bool = False) -> None:
    normalized = normalize_math_markup(text or "待补充并由发明人确认。")
    parts = re.split(r"(\$[^$]+\$)", normalized)
    has_long_formula = any(
        part.startswith("$") and part.endswith("$") and len(part) > 58
        for part in parts
    )

    if not has_long_formula:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        append_rich_text(paragraph, normalized, size=12, bold=bold, font_setter=_set_font)
        return

    pending_text = ""
    for part in parts:
        if not part:
            continue
        if part.startswith("$") and part.endswith("$") and len(part) > 58:
            pending_text = pending_text.lstrip("，,。；; ").rstrip()
            if pending_text.endswith("按"):
                pending_text = pending_text[:-1] + "依据下式："
            elif pending_text.endswith(("为", "得到")):
                pending_text += "："
            if pending_text:
                _add_body(document, pending_text, bold=bold)
            equation = document.add_paragraph()
            equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
            equation.paragraph_format.first_line_indent = Pt(0)
            equation.paragraph_format.space_before = Pt(3)
            equation.paragraph_format.space_after = Pt(3)
            equation.paragraph_format.line_spacing = 1.2
            equation._p.append(display_math_element(part[1:-1]))
            pending_text = ""
        else:
            pending_text += part
    pending_text = pending_text.lstrip("，,。；; ")
    if pending_text:
        _add_body(document, pending_text, bold=bold)


def _cn_number(index: int) -> str:
    digits = "零一二三四五六七八九"
    if index < 10:
        return digits[index]
    if index == 10:
        return "十"
    if index < 20:
        return f"十{digits[index - 10]}"
    tens, ones = divmod(index, 10)
    return f"{digits[tens]}十{digits[ones] if ones else ''}"


def _add_items(
    document: Document,
    items: list[str],
    *,
    label: str | None = None,
    numeric: bool = False,
) -> None:
    if not items:
        _add_body(document, "待补充并由发明人确认。")
        return
    for index, item in enumerate(items, start=1):
        text = item
        if label and not text.startswith(label):
            text = f"{label}{_cn_number(index)}：{text}"
        elif numeric:
            text = f"（{index}）{text}"
        _add_body(document, text)


def _format_step(text: str) -> str:
    return re.sub(r"^(S\d{3})(?:[：:,，]\s*)?", r"步骤\1，", text.strip())


def _add_patent_figures(document: Document, figures) -> int:
    inserted = 0
    for figure in figures:
        if not figure.image_path:
            continue
        image_path = Path(figure.image_path)
        if not image_path.is_file():
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run()
        with Image.open(image_path) as image:
            width_px, height_px = image.size
        width_inches = min(5.82, 7.0 * width_px / max(1, height_px))
        run.add_picture(str(image_path), width=Inches(width_inches))
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(10)
        caption.paragraph_format.keep_with_next = True
        _set_font(caption.add_run(f"图{figure.figure_no}　{figure.title}"), 10.5, bold=True)
        inserted += 1
    return inserted


def _set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_table_cell(cell, text: str, *, bold: bool = False, shaded: bool = False, left: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    _set_font(paragraph.add_run(text), 10.5, bold=bold)
    if shaded:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shd)


def _add_experiment_table(document: Document, series: list[ExperimentSeries]) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(11)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True
    _set_font(caption.add_run("表1　不同模型上的平均准确率对比（%）"), 10.5, bold=True)

    labels = series[0].labels
    table = document.add_table(rows=1 + len(labels), cols=1 + len(series))
    table.style = "Table Grid"
    widths = [3200] + [int(5800 / len(series))] * len(series)
    widths[-1] += 9000 - sum(widths)
    _set_table_geometry(table, widths)
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    _set_table_cell(header.cells[0], "干预方式", bold=True, shaded=True)
    for index, item in enumerate(series, start=1):
        _set_table_cell(header.cells[index], item.name, bold=True, shaded=True)
    best_values = [max(item.values) for item in series]
    for row_index, label in enumerate(labels, start=1):
        _set_table_cell(table.rows[row_index].cells[0], label, left=True)
        for column_index, item in enumerate(series, start=1):
            value = item.values[row_index - 1]
            _set_table_cell(
                table.rows[row_index].cells[column_index],
                f"{value:.1f}",
                bold=value == best_values[column_index - 1],
            )


def _add_experiment_chart(document: Document, image_path: Path, figure_no: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(image_path), width=Inches(5.82))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    _set_font(
        caption.add_run(f"图{figure_no}　不同向量来源与适配方式的平均准确率对比图"),
        10.5,
        bold=True,
    )


def docx(
    job: PipelineJob,
    output_path: Path,
    *,
    experiment_series: list[ExperimentSeries] | None = None,
    experiment_chart_path: Path | None = None,
    experiment_figure_no: int | None = None,
) -> None:
    disclosure = job.disclosure
    if disclosure is None:
        raise ValueError("技术交底书尚未生成")
    document = Document()
    section = document.sections[0]
    # Chinese technical disclosures are normally exchanged on ISO A4 paper.
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(20)
    _set_font(title.add_run("技　术　交　底　书"), 18, bold=True, east_asia_font=HEADING_CN_FONT)

    _add_heading(document, "1. 发明创造名称", 1)
    _add_body(document, disclosure.invention_title, bold=True)

    _add_heading(document, "2. 背景技术及现有技术的缺陷和不足", 1)
    _add_heading(document, "2.1 背景技术", 2)
    _add_heading(document, "（一）技术领域", 3)
    _add_body(document, disclosure.technical_field)
    _add_heading(document, "（二）相关背景技术", 3)
    _add_items(document, disclosure.background)
    _add_heading(document, "2.2 现有技术的缺陷和不足", 2)
    _add_items(document, disclosure.prior_art_defects, label="缺陷")
    _add_body(document, "针对上述缺陷，本发明所要解决的技术问题是：")
    _add_items(document, disclosure.technical_problem, numeric=True)

    figures = list(job.evidence_package.patent_figures) if job.evidence_package else []
    _add_heading(document, "3. 具体的技术方案描述", 1)
    _add_heading(document, "3.1 总体构思", 2)
    _add_body(document, disclosure.overall_solution)
    _add_heading(document, "3.2 详细技术方案", 2)
    for step in disclosure.detailed_steps:
        _add_body(document, _format_step(step))
    if figures:
        _add_patent_figures(document, figures[:1])

    _add_heading(document, "3.3 系统实现与数据接口", 2)
    _add_items(document, disclosure.system_implementation)
    _add_items(document, disclosure.data_and_interfaces)
    if len(figures) > 1:
        _add_patent_figures(document, figures[1:2])

    _add_heading(document, "3.4 发明点归纳", 2)
    _add_items(document, disclosure.key_innovations, label="发明点")
    _add_heading(document, "3.5 术语、实施边界与替代方案", 2)
    _add_items(document, disclosure.terminology)
    _add_items(document, disclosure.implementation_boundaries)
    _add_items(document, disclosure.alternatives)

    _add_heading(document, "4. 本发明创造的优点", 1)
    _add_items(document, disclosure.beneficial_effects, label="优点")

    _add_heading(document, "5. 具体实施方式及附图", 1)
    subsection = 1
    if disclosure.embodiments:
        for index, embodiment in enumerate(disclosure.embodiments, start=1):
            _add_heading(document, f"5.{subsection} 实施例{_cn_number(index)}", 2)
            _add_body(document, embodiment)
            subsection += 1
    else:
        _add_heading(document, f"5.{subsection} 具体实施例", 2)
        _add_body(document, "待补充并由发明人确认。")
        subsection += 1

    _add_heading(document, f"5.{subsection} 实验验证", 2)
    _add_items(document, disclosure.experimental_evidence)
    subsection += 1
    if experiment_series and experiment_chart_path is not None and experiment_figure_no is not None:
        document.add_page_break()
        _add_heading(document, f"5.{subsection} 实验数据图表", 2)
        _add_experiment_table(document, experiment_series)
        _add_experiment_chart(document, experiment_chart_path, experiment_figure_no)
        subsection += 1
    _add_heading(document, f"5.{subsection} 附图说明", 2)
    _add_items(document, disclosure.drawing_descriptions)
    if len(figures) > 2:
        _add_patent_figures(document, figures[2:])
    subsection += 1
    _add_heading(document, f"5.{subsection} 发明人确认事项", 2)
    _add_items(document, disclosure.inventor_confirmation_items, numeric=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def export_job(job: PipelineJob, data_dir: Path) -> dict[str, str]:
    if job.evidence_package is None or job.solution is None:
        raise ValueError("缺少专利附图所需的技术方案或证据包")
    root = data_dir / job.job_id / "artifacts"
    root.mkdir(parents=True, exist_ok=True)
    markdown_path = root / "technical_disclosure.md"
    docx_path = root / "technical_disclosure.docx"
    json_path = root / "job.json"
    figures_dir = root / "figures"
    figures_zip_path = root / "patent_figures.zip"
    figures = job.evidence_package.patent_figures or build_default_patent_figures(job.solution)
    job.evidence_package.patent_figures = render_patent_figures(figures, figures_dir)
    descriptions = drawing_descriptions(job.evidence_package.patent_figures)
    experiment_series = extract_experiment_series(job.disclosure.experimental_evidence if job.disclosure else [])
    experiment_figure_no = len(job.evidence_package.patent_figures) + 1
    experiment_chart_path = render_experiment_chart(
        experiment_series,
        figures_dir,
        experiment_figure_no,
    )
    if experiment_chart_path is not None:
        descriptions.append(f"图{experiment_figure_no}为不同向量来源与适配方式的平均准确率对比图。")
    job.evidence_package.figure_plan = descriptions
    if job.disclosure is not None:
        job.disclosure.drawing_descriptions = descriptions
    artifacts = {
        "markdown": str(markdown_path.resolve()),
        "docx": str(docx_path.resolve()),
        "json": str(json_path.resolve()),
        "figures": str(figures_zip_path.resolve()),
    }
    job.artifacts = artifacts
    markdown_path.write_text(markdown(job), encoding="utf-8")
    docx(
        job,
        docx_path,
        experiment_series=experiment_series,
        experiment_chart_path=experiment_chart_path,
        experiment_figure_no=experiment_figure_no if experiment_chart_path is not None else None,
    )
    with zipfile.ZipFile(figures_zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(figures_dir.iterdir()):
            if path.is_file():
                archive.write(path, arcname=path.name)
    return artifacts
