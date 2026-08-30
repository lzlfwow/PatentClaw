from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from ..constants import DISCLOSURE_SECTION_LABELS
from ..schemas import PatentFigureAsset, TechnicalDisclosure


def write_disclosure_docx(
    disclosure: TechnicalDisclosure,
    path: Path,
    figures: list[PatentFigureAsset] | None = None,
) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    title = document.add_heading(disclosure.invention_title or "技术交底书", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for field, label in DISCLOSURE_SECTION_LABELS.items():
        if field == "invention_title":
            continue
        document.add_heading(label, level=1)
        value = getattr(disclosure, field)
        if isinstance(value, list):
            if value:
                for item in value:
                    document.add_paragraph(item)
            else:
                document.add_paragraph("（待补充）")
        else:
            document.add_paragraph(value or "（待补充）")
    if figures:
        document.add_heading("专利附图", level=1)
        for figure in figures:
            image_path = Path(figure.image_path) if figure.image_path else None
            if image_path is not None and image_path.is_file():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
            caption = document.add_paragraph(f"图{figure.figure_no}　{figure.title}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save(path)
